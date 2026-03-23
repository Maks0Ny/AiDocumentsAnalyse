from dataclasses import dataclass
from datetime import datetime
from typing import List, Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


@dataclass
class ReportMetadata:
    contract_name: str
    source_path: str
    generated_at: str
    analysis_type: str = "Гибридный анализ (правила + нейросеть)"


class ContractReportGenerator:
    """
    Генератор итогового отчёта по договору.
    Поддерживает:
    1. DOCX-отчёт
    2. TXT-отчёт
    """

    def generate_docx_report(
        self,
        output_path: str,
        metadata: ReportMetadata,
        final_risks: List,
        summary: Optional[Dict[str, int]] = None,
    ) -> None:
        if Document is None:
            raise ImportError(
                "Библиотека python-docx не установлена. "
                "Установите её командой: pip install python-docx"
            )

        doc = Document()
        self._set_default_styles(doc)

        # Заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("ИТОГОВЫЙ ОТЧЁТ\nПО АНАЛИЗУ ДОГОВОРА")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph("")

        # Общая информация
        doc.add_heading("1. Общая информация", level=1)
        doc.add_paragraph(f"Наименование документа: {metadata.contract_name}")
        doc.add_paragraph(f"Исходный файл: {metadata.source_path}")
        doc.add_paragraph(f"Дата и время формирования отчёта: {metadata.generated_at}")
        doc.add_paragraph(f"Тип анализа: {metadata.analysis_type}")

        # Сводка
        doc.add_heading("2. Сводка результатов анализа", level=1)
        if summary:
            doc.add_paragraph(f"Общее количество выявленных рисков: {summary.get('total_risks', 0)}")
            doc.add_paragraph(f"Высокий риск: {summary.get('high_risk', 0)}")
            doc.add_paragraph(f"Средний риск: {summary.get('medium_risk', 0)}")
            doc.add_paragraph(f"Низкий риск: {summary.get('low_risk', 0)}")
            doc.add_paragraph(f"Количество категорий рисков: {summary.get('categories_count', 0)}")
        else:
            doc.add_paragraph(f"Общее количество выявленных рисков: {len(final_risks)}")

        # Общее заключение
        doc.add_heading("3. Общее заключение", level=1)
        conclusion_text = self._build_conclusion(final_risks, summary)
        doc.add_paragraph(conclusion_text)

        # Таблица рисков
        doc.add_heading("4. Перечень выявленных рисков", level=1)
        if final_risks:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "№"
            hdr_cells[1].text = "Категория"
            hdr_cells[2].text = "Уровень риска"
            hdr_cells[3].text = "Название"
            hdr_cells[4].text = "Фрагмент"

            for idx, risk in enumerate(final_risks, start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = getattr(risk, "category", "")
                row_cells[2].text = getattr(risk, "risk_level", "")
                row_cells[3].text = getattr(risk, "title", "")
                row_cells[4].text = getattr(risk, "fragment_text", "")
        else:
            doc.add_paragraph("По результатам анализа риски не обнаружены.")

        # Подробное описание рисков
        doc.add_heading("5. Подробный анализ и рекомендации", level=1)
        if final_risks:
            grouped = self._group_risks_by_category(final_risks)

            for category, risks in grouped.items():
                doc.add_heading(category, level=2)

                for idx, risk in enumerate(risks, start=1):
                    p = doc.add_paragraph()
                    p.add_run(f"{idx}. {getattr(risk, 'title', '')}").bold = True

                    doc.add_paragraph(f"Уровень риска: {getattr(risk, 'risk_level', '')}")
                    doc.add_paragraph(f"Описание: {getattr(risk, 'description', '')}")

                    fragment_paragraph = doc.add_paragraph()
                    fragment_paragraph.add_run("Проблемный фрагмент: ").bold = True
                    fragment_paragraph.add_run(getattr(risk, "fragment_text", ""))

                    recommendation_paragraph = doc.add_paragraph()
                    recommendation_paragraph.add_run("Рекомендация: ").bold = True
                    recommendation_paragraph.add_run(getattr(risk, "recommendation", ""))

                    doc.add_paragraph("")
        else:
            doc.add_paragraph("Подробный анализ отсутствует, так как риски не были выявлены.")

        # Итоговые рекомендации
        doc.add_heading("6. Итоговые рекомендации", level=1)
        final_recommendations = self._build_final_recommendations(final_risks)
        if final_recommendations:
            for rec in final_recommendations:
                doc.add_paragraph(rec, style="List Bullet")
        else:
            doc.add_paragraph("Существенных замечаний по результатам анализа не выявлено.")

        doc.save(output_path)

    def generate_txt_report(
        self,
        output_path: str,
        metadata: ReportMetadata,
        final_risks: List,
        summary: Optional[Dict[str, int]] = None,
    ) -> None:
        lines = []
        lines.append("ИТОГОВЫЙ ОТЧЁТ ПО АНАЛИЗУ ДОГОВОРА")
        lines.append("=" * 80)
        lines.append("")

        lines.append("1. ОБЩАЯ ИНФОРМАЦИЯ")
        lines.append(f"Наименование документа: {metadata.contract_name}")
        lines.append(f"Исходный файл: {metadata.source_path}")
        lines.append(f"Дата и время формирования отчёта: {metadata.generated_at}")
        lines.append(f"Тип анализа: {metadata.analysis_type}")
        lines.append("")

        lines.append("2. СВОДКА")
        if summary:
            lines.append(f"Общее количество выявленных рисков: {summary.get('total_risks', 0)}")
            lines.append(f"Высокий риск: {summary.get('high_risk', 0)}")
            lines.append(f"Средний риск: {summary.get('medium_risk', 0)}")
            lines.append(f"Низкий риск: {summary.get('low_risk', 0)}")
            lines.append(f"Количество категорий рисков: {summary.get('categories_count', 0)}")
        else:
            lines.append(f"Общее количество выявленных рисков: {len(final_risks)}")
        lines.append("")

        lines.append("3. ОБЩЕЕ ЗАКЛЮЧЕНИЕ")
        lines.append(self._build_conclusion(final_risks, summary))
        lines.append("")

        lines.append("4. ПОДРОБНЫЙ ПЕРЕЧЕНЬ РИСКОВ")
        if final_risks:
            grouped = self._group_risks_by_category(final_risks)
            for category, risks in grouped.items():
                lines.append(f"\n[{category}]")
                for idx, risk in enumerate(risks, start=1):
                    lines.append(f"{idx}. {getattr(risk, 'title', '')}")
                    lines.append(f"   Уровень риска: {getattr(risk, 'risk_level', '')}")
                    lines.append(f"   Описание: {getattr(risk, 'description', '')}")
                    lines.append(f"   Фрагмент: {getattr(risk, 'fragment_text', '')}")
                    lines.append(f"   Рекомендация: {getattr(risk, 'recommendation', '')}")
                    lines.append("")
        else:
            lines.append("Риски не обнаружены.")
        lines.append("")

        lines.append("5. ИТОГОВЫЕ РЕКОМЕНДАЦИИ")
        recommendations = self._build_final_recommendations(final_risks)
        if recommendations:
            for rec in recommendations:
                lines.append(f"- {rec}")
        else:
            lines.append("Существенных замечаний не выявлено.")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _set_default_styles(self, doc) -> None:
        styles = doc.styles
        if "Normal" in styles:
            normal_style = styles["Normal"]
            normal_style.font.name = "Times New Roman"
            normal_style.font.size = Pt(12)

    def _group_risks_by_category(self, risks: List) -> Dict[str, List]:
        grouped: Dict[str, List] = {}
        for risk in risks:
            category = getattr(risk, "category", "Без категории")
            grouped.setdefault(category, []).append(risk)
        return grouped

    def _build_conclusion(self, final_risks: List, summary: Optional[Dict[str, int]]) -> str:
        if not final_risks:
            return (
                "По результатам автоматизированного анализа существенные юридические риски "
                "в тексте договора не обнаружены. Документ рекомендуется дополнительно "
                "проверить специалистом для окончательной правовой оценки."
            )

        high_risk = summary.get("high_risk", 0) if summary else 0
        medium_risk = summary.get("medium_risk", 0) if summary else 0
        low_risk = summary.get("low_risk", 0) if summary else 0

        if high_risk > 0:
            return (
                "В результате анализа в договоре выявлены существенные юридические риски, "
                "включая положения высокого уровня значимости. Документ требует доработки "
                "до его подписания и практического применения."
            )
        if medium_risk > 0:
            return (
                "В результате анализа в договоре выявлены риски среднего уровня значимости. "
                "Рекомендуется уточнить отдельные формулировки и дополнить недостающие положения."
            )
        if low_risk > 0:
            return (
                "В тексте договора выявлены отдельные риски низкого уровня значимости, "
                "связанные преимущественно с неопределённостью формулировок и структурой документа."
            )

        return (
            "По результатам анализа обнаружены замечания, требующие дополнительного рассмотрения."
        )

    def _build_final_recommendations(self, final_risks: List) -> List[str]:
        recommendations = []
        seen = set()

        for risk in final_risks:
            rec = getattr(risk, "recommendation", "").strip()
            if rec and rec.lower() not in seen:
                seen.add(rec.lower())
                recommendations.append(rec)

        return recommendations


def build_report_metadata(contract_name: str, source_path: str) -> ReportMetadata:
    return ReportMetadata(
        contract_name=contract_name,
        source_path=source_path,
        generated_at=datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
    )